"""
Testes de Entrega de Materiais (Termo de Entrega de Materiais) e do Termo
de Responsabilidade — ambos exportados em .docx, no layout oficial do
modelo. Cobre:
- RBAC: exclusiva do MASTER — GESTOR_POLO não cadastra/exporta entregas,
  nem mesmo do próprio polo.
- A entrega nasce com o coordenador copiado do responsável do polo.
- O termo de entrega e o termo de responsabilidade trazem os dados
  cadastrados (do polo/entrega) nos parágrafos e células certas.
"""
import io

import docx

from tests.conftest import login


def test_gestor_de_polo_nao_cria_entrega(client, seed_basico):
    token = login(client, "gestor.a@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    resp = client.post(
        "/api/v1/entregas-materiais",
        json={"polo_id": polo_a_id, "itens": []},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 403


def test_entrega_nasce_com_coordenador_do_polo(client, seed_basico):
    token_master = login(client, "master@test.com")
    headers_master = {"Authorization": f"Bearer {token_master}"}
    polo_a_id = str(seed_basico["polo_a"].id)
    client.patch(f"/api/v1/polos/{polo_a_id}", json={"responsavel_nome": "Coordenadora Fulana"}, headers=headers_master)

    resp = client.post(
        "/api/v1/entregas-materiais",
        json={
            "polo_id": polo_a_id, "data_entrega": "2026-03-01",
            "itens": [{"descricao": "Bolas de futebol", "quantidade": "10"}],
        },
        headers=headers_master,
    )
    assert resp.status_code == 201, resp.text
    assert resp.json()["coordenador_nome"] == "Coordenadora Fulana"


def test_exportar_termo_entrega_reflete_itens_e_coordenador(client, seed_basico):
    token_master = login(client, "master@test.com")
    headers_master = {"Authorization": f"Bearer {token_master}"}
    polo_a_id = str(seed_basico["polo_a"].id)
    client.patch(f"/api/v1/polos/{polo_a_id}", json={"responsavel_nome": "Coordenadora Fulana"}, headers=headers_master)

    headers = headers_master
    entrega = client.post(
        "/api/v1/entregas-materiais",
        json={
            "polo_id": polo_a_id, "data_entrega": "2026-03-01",
            "itens": [
                {"descricao": "Bolas de futebol", "quantidade": "10"},
                {"descricao": "Coletes", "quantidade": "20"},
            ],
        },
        headers=headers,
    ).json()

    resp = client.get(f"/api/v1/entregas-materiais/{entrega['id']}/exportar", headers=headers)
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    doc = docx.Document(io.BytesIO(resp.content))
    textos = [p.text for p in doc.paragraphs]
    assert any(t == "NÚCLEO: Polo A" for t in textos)
    assert any(t == "COORDENADOR: Coordenadora Fulana" for t in textos)

    tabela = doc.tables[0]
    assert tabela.cell(1, 0).text == "Bolas de futebol"
    assert tabela.cell(1, 1).text == "10"
    assert tabela.cell(2, 0).text == "Coletes"
    assert tabela.cell(2, 1).text == "20"


def test_exportar_termo_responsabilidade_reflete_dados_do_polo(client, seed_basico):
    token = login(client, "master@test.com")
    headers = {"Authorization": f"Bearer {token}"}
    polo_a_id = str(seed_basico["polo_a"].id)
    client.patch(
        f"/api/v1/polos/{polo_a_id}",
        json={
            "representante_legal_nome": "Fulana de Tal",
            "representante_legal_cpf": "111.222.333-44",
            "representante_legal_rg": "12.345.678-9",
            "representante_legal_endereco": "Rua das Flores, 100",
            "representante_legal_bairro": "Centro",
            "representante_legal_cidade": "São Paulo",
        },
        headers=headers,
    )

    resp = client.get(f"/api/v1/polos/{polo_a_id}/termo-responsabilidade/exportar", headers=headers)
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    doc = docx.Document(io.BytesIO(resp.content))
    texto_completo = "\n".join(p.text for p in doc.paragraphs)
    assert "Fulana de Tal" in texto_completo
    assert "RG nº 12.345.678-9" in texto_completo
    assert "CPF nº 111.222.333-44" in texto_completo
    assert "Rua das Flores, 100" in texto_completo
    assert "Bairro Centro" in texto_completo
    assert "Cidade São Paulo" in texto_completo
    assert any(p.text.startswith("São Paulo,") for p in doc.paragraphs)
