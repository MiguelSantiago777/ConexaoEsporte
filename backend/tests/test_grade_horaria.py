"""
Testes da exportação da Grade Horária (por polo) em .docx, no layout
oficial do modelo. Cobre:
- RBAC: MASTER e GESTOR_POLO do próprio polo podem exportar; gestor de
  outro polo não pode.
- A grade reflete as turmas do polo (letra, dias/carga horária) e soma
  corretamente o total semanal + planejamento.
"""
import io

import docx

from tests.conftest import login


def _criar_turma(client, token, polo_id, modalidade_id, horario_inicio, horario_fim, dias_semana):
    return client.post(
        "/api/v1/turmas",
        json={
            "polo_id": polo_id, "modalidade_id": modalidade_id,
            "horario_inicio": horario_inicio, "horario_fim": horario_fim,
            "dias_semana": dias_semana, "limite_vagas": 10,
        },
        headers={"Authorization": f"Bearer {token}"},
    ).json()


def test_gestor_de_outro_polo_nao_exporta_grade_horaria(client, seed_basico):
    token = login(client, "gestor.b@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    resp = client.get(
        f"/api/v1/polos/{polo_a_id}/grade-horaria/exportar",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 403


def test_exportar_grade_horaria_reflete_turmas_do_polo(client, seed_basico):
    token = login(client, "gestor.a@test.com")
    headers = {"Authorization": f"Bearer {token}"}
    polo_a_id = str(seed_basico["polo_a"].id)
    modalidade_id = str(seed_basico["modalidade"].id)

    natacao = client.post(
        "/api/v1/modalidades", json={"nome": "Natação"}, headers=headers
    ).json()

    _criar_turma(client, token, polo_a_id, modalidade_id, "08:00", "09:30", ["SEG", "QUA", "SEX"])
    _criar_turma(client, token, polo_a_id, natacao["id"], "10:00", "11:00", ["QUA"])

    resp = client.get(
        f"/api/v1/polos/{polo_a_id}/grade-horaria/exportar",
        params={"planejamento_horas": 2},
        headers=headers,
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    doc = docx.Document(io.BytesIO(resp.content))
    tabela = doc.tables[0]

    assert tabela.cell(0, 0).text == "HORAS AULAS SEMANAL NÚCLEO - Polo A"
    # turma A: Judô, 08:00-09:30 (1.5h), Seg/Qua/Sex
    assert tabela.cell(2, 0).text == "A"
    assert tabela.cell(2, 1).text == "1.5"  # Segunda
    assert tabela.cell(2, 2).text == "1.5"  # Quarta
    assert tabela.cell(2, 3).text == "1.5"  # Sexta
    # turma B: Natação, 10:00-11:00 (1h), só Quarta
    assert tabela.cell(3, 0).text == "B"
    assert tabela.cell(3, 1).text == ""     # Segunda (não tem aula)
    assert tabela.cell(3, 2).text == "1"    # Quarta
    assert tabela.cell(3, 3).text == ""     # Sexta

    # total semanal = 1.5*3 + 1 = 5.5h ; planejamento = 2h ; total geral = 7.5h
    assert "5.5" in tabela.cell(9, 4).text
    assert "2" in tabela.cell(10, 4).text
    assert "7.5" in tabela.cell(11, 4).text

    texto_completo = "\n".join(p.text for p in doc.paragraphs)
    assert "Judô" in texto_completo
    assert "Natação" in texto_completo
