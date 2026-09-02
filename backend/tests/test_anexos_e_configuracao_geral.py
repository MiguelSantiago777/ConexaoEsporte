"""
Testes dos anexos do cadastro de professor (foto/documento/contrato), do
repositório de Anexos Gerais por polo e da Configuração Geral (nome do
projeto, número de convênio e vigência, exibidos no cabeçalho dos
relatórios, junto do logo do projeto).
"""
import shutil
from pathlib import Path

import pytest

from tests.conftest import login


@pytest.fixture(autouse=True)
def _limpar_uploads_teste():
    yield
    shutil.rmtree(Path("uploads/usuarios"), ignore_errors=True)
    shutil.rmtree(Path("uploads/anexos_gerais"), ignore_errors=True)


def _criar_professor(client, token_gestor, polo_id, email="prof.anexo@test.com"):
    resp = client.post(
        "/api/v1/usuarios",
        json={"nome": "Professor Teste", "email": email, "senha": "senha123", "perfil": "PROFESSOR", "polo_id": polo_id},
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp.status_code == 201, resp.text
    return resp.json()


def test_gestor_anexa_lista_baixa_e_remove_documento_do_professor(client, seed_basico):
    token_gestor = login(client, "gestor.a@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    professor = _criar_professor(client, token_gestor, polo_a_id)

    resp = client.post(
        f"/api/v1/usuarios/{professor['id']}/documentos",
        data={"tipo": "FOTO"},
        files={"arquivo": ("foto.jpg", b"foto-fake", "image/jpeg")},
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp.status_code == 201, resp.text
    documento = resp.json()
    assert documento["tipo"] == "FOTO"
    assert documento["nome_arquivo"] == "foto.jpg"

    resp_lista = client.get(
        f"/api/v1/usuarios/{professor['id']}/documentos",
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp_lista.status_code == 200
    assert len(resp_lista.json()) == 1

    resp_arquivo = client.get(
        f"/api/v1/usuarios/documentos/{documento['id']}/arquivo",
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp_arquivo.status_code == 200
    assert resp_arquivo.content == b"foto-fake"

    resp_del = client.delete(
        f"/api/v1/usuarios/documentos/{documento['id']}",
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp_del.status_code == 204

    resp_lista2 = client.get(
        f"/api/v1/usuarios/{professor['id']}/documentos",
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp_lista2.json() == []


def test_gestor_de_outro_polo_nao_acessa_documentos_do_professor(client, seed_basico):
    token_gestor_a = login(client, "gestor.a@test.com")
    token_gestor_b = login(client, "gestor.b@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    professor = _criar_professor(client, token_gestor_a, polo_a_id)

    resp = client.post(
        f"/api/v1/usuarios/{professor['id']}/documentos",
        data={"tipo": "CONTRATO"},
        files={"arquivo": ("contrato.pdf", b"conteudo", "application/pdf")},
        headers={"Authorization": f"Bearer {token_gestor_b}"},
    )
    assert resp.status_code == 403


def test_anexo_geral_por_polo_e_isolado_entre_polos(client, seed_basico):
    token_gestor_a = login(client, "gestor.a@test.com")
    token_gestor_b = login(client, "gestor.b@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    polo_b_id = str(seed_basico["polo_b"].id)

    resp = client.post(
        "/api/v1/anexos-gerais",
        data={"polo_id": polo_a_id, "titulo": "Apólice de seguro"},
        files={"arquivo": ("apolice.pdf", b"conteudo", "application/pdf")},
        headers={"Authorization": f"Bearer {token_gestor_a}"},
    )
    assert resp.status_code == 201, resp.text
    anexo = resp.json()

    # Gestor do polo B não enxerga o anexo do polo A.
    resp_lista_b = client.get(
        "/api/v1/anexos-gerais", headers={"Authorization": f"Bearer {token_gestor_b}"}
    )
    assert resp_lista_b.json() == []

    # Nem consegue baixá-lo ou removê-lo diretamente pelo id.
    resp_arquivo_b = client.get(
        f"/api/v1/anexos-gerais/{anexo['id']}/arquivo", headers={"Authorization": f"Bearer {token_gestor_b}"}
    )
    assert resp_arquivo_b.status_code == 403

    resp_lista_a = client.get(
        "/api/v1/anexos-gerais", headers={"Authorization": f"Bearer {token_gestor_a}"}
    )
    assert len(resp_lista_a.json()) == 1

    resp_del = client.delete(
        f"/api/v1/anexos-gerais/{anexo['id']}", headers={"Authorization": f"Bearer {token_gestor_a}"}
    )
    assert resp_del.status_code == 204


def _criar_turma_consolidado(client, token, polo_id, modalidade_id, professor_id=None):
    turma = client.post(
        "/api/v1/turmas",
        json={
            "polo_id": polo_id, "modalidade_id": modalidade_id,
            "horario_inicio": "08:00", "horario_fim": "09:00",
            "dias_semana": ["SEG"], "limite_vagas": 10,
        },
        headers={"Authorization": f"Bearer {token}"},
    ).json()
    if professor_id:
        client.patch(
            f"/api/v1/turmas/{turma['id']}", json={"professor_id": professor_id},
            headers={"Authorization": f"Bearer {token}"},
        )
    return turma


def test_visao_consolidada_reune_anexo_geral_foto_de_chamada_e_observacao_de_aula(client, seed_basico):
    token_master = login(client, "master@test.com")
    token_gestor_a = login(client, "gestor.a@test.com")
    polo_a_id = str(seed_basico["polo_a"].id)
    polo_b_id = str(seed_basico["polo_b"].id)
    modalidade_id = str(seed_basico["modalidade"].id)

    professor = _criar_professor(client, token_gestor_a, polo_a_id, "prof.consolidado@test.com")
    turma = _criar_turma_consolidado(client, token_gestor_a, polo_a_id, modalidade_id, professor_id=professor["id"])
    token_prof = login(client, professor["email"])

    resp_anexo = client.post(
        "/api/v1/anexos-gerais",
        data={"polo_id": polo_a_id, "titulo": "Apólice de seguro"},
        files={"arquivo": ("apolice.pdf", b"conteudo", "application/pdf")},
        headers={"Authorization": f"Bearer {token_gestor_a}"},
    )
    assert resp_anexo.status_code == 201, resp_anexo.text

    resp_evidencia = client.post(
        "/api/v1/frequencias/evidencias",
        data={"turma_id": turma["id"], "data": "2026-03-02"},
        files={"arquivos": ("chamada.jpg", b"foto-fake", "image/jpeg")},
        headers={"Authorization": f"Bearer {token_prof}"},
    )
    assert resp_evidencia.status_code == 201, resp_evidencia.text

    resp_relatorio = client.post(
        "/api/v1/relatorios-aula",
        json={"turma_id": turma["id"], "data": "2026-03-02", "conteudo_trabalhado": "Fundamentos", "observacoes": "Turma dispersa hoje."},
        headers={"Authorization": f"Bearer {token_prof}"},
    )
    assert resp_relatorio.status_code == 201, resp_relatorio.text

    # Relatório sem observações não deve aparecer na listagem.
    resp_relatorio_sem_obs = client.post(
        "/api/v1/relatorios-aula",
        json={"turma_id": turma["id"], "data": "2026-03-09", "conteudo_trabalhado": "Fundamentos 2"},
        headers={"Authorization": f"Bearer {token_prof}"},
    )
    assert resp_relatorio_sem_obs.status_code == 201, resp_relatorio_sem_obs.text

    resp_consolidado = client.get(
        "/api/v1/anexos-gerais/consolidado", headers={"Authorization": f"Bearer {token_master}"}
    )
    assert resp_consolidado.status_code == 200, resp_consolidado.text
    itens = resp_consolidado.json()
    tipos = {item["tipo"] for item in itens}
    assert tipos == {"ANEXO_GERAL", "EVIDENCIA_CHAMADA", "OBSERVACAO_AULA"}
    assert len(itens) == 3

    observacao = next(i for i in itens if i["tipo"] == "OBSERVACAO_AULA")
    assert observacao["descricao"] == "Turma dispersa hoje."
    assert observacao["autor_nome"] == "Professor Teste"
    assert observacao["polo_nome"] == "Polo A"

    evidencia = next(i for i in itens if i["tipo"] == "EVIDENCIA_CHAMADA")
    assert evidencia["nome_arquivo"] == "chamada.jpg"
    assert evidencia["possui_arquivo"] is True

    # Gestor do polo B só enxerga o que é do próprio polo (nada, aqui).
    token_gestor_b = login(client, "gestor.b@test.com")
    resp_gestor_b = client.get(
        "/api/v1/anexos-gerais/consolidado", headers={"Authorization": f"Bearer {token_gestor_b}"}
    )
    assert resp_gestor_b.json() == []

    # MASTER filtrando por polo_id do polo B (sem dados) fica vazio; polo A traz os 3 itens.
    resp_master_polo_b = client.get(
        "/api/v1/anexos-gerais/consolidado", params={"polo_id": polo_b_id},
        headers={"Authorization": f"Bearer {token_master}"},
    )
    assert resp_master_polo_b.json() == []


def test_configuracao_geral_e_exclusiva_do_master(client, seed_basico):
    token_master = login(client, "master@test.com")
    token_gestor = login(client, "gestor.a@test.com")

    resp_get_vazio = client.get("/api/v1/configuracao-geral", headers={"Authorization": f"Bearer {token_master}"})
    assert resp_get_vazio.status_code == 200
    assert resp_get_vazio.json() is None

    resp_negado = client.patch(
        "/api/v1/configuracao-geral",
        json={"numero_convenio": "123/2026"},
        headers={"Authorization": f"Bearer {token_gestor}"},
    )
    assert resp_negado.status_code == 403

    resp_patch = client.patch(
        "/api/v1/configuracao-geral",
        json={
            "nome_projeto": "Conexão Esporte",
            "numero_convenio": "123/2026",
            "data_inicio_projeto": "2026-01-01",
            "data_fim_projeto": "2026-12-31",
        },
        headers={"Authorization": f"Bearer {token_master}"},
    )
    assert resp_patch.status_code == 200, resp_patch.text
    assert resp_patch.json()["numero_convenio"] == "123/2026"
    assert resp_patch.json()["nome_projeto"] == "Conexão Esporte"

    resp_get = client.get("/api/v1/configuracao-geral", headers={"Authorization": f"Bearer {token_master}"})
    assert resp_get.json()["numero_convenio"] == "123/2026"
    assert resp_get.json()["nome_projeto"] == "Conexão Esporte"

    # Editar de novo (idempotente — continua sendo o mesmo registro único).
    resp_patch2 = client.patch(
        "/api/v1/configuracao-geral",
        json={"numero_convenio": "456/2027", "data_inicio_projeto": None, "data_fim_projeto": None},
        headers={"Authorization": f"Bearer {token_master}"},
    )
    assert resp_patch2.json()["numero_convenio"] == "456/2027"
    assert resp_patch2.json()["data_inicio_projeto"] is None


def test_texto_cabecalho_omite_partes_ausentes():
    from app.application.relatorios.cabecalho_convenio import texto_cabecalho
    from app.domain.configuracao_geral.entities import ConfiguracaoGeral
    from datetime import date

    assert texto_cabecalho(None) is None
    assert texto_cabecalho(
        ConfiguracaoGeral(id=None, nome_projeto=None, numero_convenio=None, data_inicio_projeto=None, data_fim_projeto=None)
    ) is None

    so_numero = texto_cabecalho(
        ConfiguracaoGeral(
            id=None, nome_projeto=None, numero_convenio="123/2026", data_inicio_projeto=None, data_fim_projeto=None
        )
    )
    assert so_numero == "Convênio nº 123/2026"

    completo = texto_cabecalho(
        ConfiguracaoGeral(
            id=None, nome_projeto="Conexão Esporte", numero_convenio="123/2026",
            data_inicio_projeto=date(2026, 1, 1), data_fim_projeto=date(2026, 12, 31),
        )
    )
    assert completo == "Projeto: Conexão Esporte — Convênio nº 123/2026 — Vigência do projeto: 01/01/2026 a 31/12/2026"


def test_cabecalho_aplicado_no_termo_de_entrega_gerado():
    import docx
    import io

    from app.application.relatorios.termo_entrega_export_service import exportar_termo_entrega

    buffer = exportar_termo_entrega(
        polo_nome="Polo A", coordenador_nome="Fulano", itens=[], cabecalho_convenio="Convênio nº 999"
    )
    documento = docx.Document(io.BytesIO(buffer.getvalue()))
    header = documento.sections[0].header
    texto_cabecalho_gerado = " ".join(p.text for p in header.paragraphs)
    assert "Convênio nº 999" in texto_cabecalho_gerado
    # logo do projeto embutido no cabeçalho (pelo menos uma imagem inline)
    assert any(run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
               for p in header.paragraphs for run in p.runs)
