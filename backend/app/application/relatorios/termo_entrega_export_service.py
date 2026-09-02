"""
Exportação do Termo de Entrega de Materiais para .docx, no layout oficial
do modelo. Carrega o arquivo-modelo em
`app/infrastructure/templates/termo_entrega_materiais.docx` — os rótulos
"NÚCLEO:" e "COORDENADOR:" não têm uma célula de valor separada, então o
próprio texto do parágrafo é sobrescrito com "rótulo + valor", mesmo padrão
já usado nos demais exportadores.

Limite do modelo oficial: até 18 itens na tabela — excedentes não entram
na exportação.
"""
import io
from dataclasses import dataclass
from pathlib import Path

import docx

from app.application.relatorios.cabecalho_convenio import aplicar_cabecalho_docx

TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "infrastructure" / "templates" / "termo_entrega_materiais.docx"

MAX_ITENS = 18


@dataclass
class ItemEntrega:
    descricao: str
    quantidade: str


def _substituir_paragrafo(paragrafo, texto: str) -> None:
    for run in paragrafo.runs:
        run.text = ""
    if paragrafo.runs:
        paragrafo.runs[0].text = texto
    else:
        paragrafo.add_run(texto)


def exportar_termo_entrega(
    *, polo_nome: str, coordenador_nome: str, itens: list[ItemEntrega], cabecalho_convenio: str | None = None
) -> io.BytesIO:
    wb = docx.Document(TEMPLATE_PATH)

    for paragrafo in wb.paragraphs:
        if paragrafo.text.startswith("NÚCLEO:"):
            _substituir_paragrafo(paragrafo, f"NÚCLEO: {polo_nome}")
        elif paragrafo.text.startswith("COORDENADOR:"):
            _substituir_paragrafo(paragrafo, f"COORDENADOR: {coordenador_nome}")

    tabela = wb.tables[0]
    for idx, item in enumerate(itens[:MAX_ITENS]):
        linha = 1 + idx
        tabela.cell(linha, 0).text = item.descricao
        tabela.cell(linha, 1).text = item.quantidade

    aplicar_cabecalho_docx(wb, cabecalho_convenio)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
