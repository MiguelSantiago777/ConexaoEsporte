"""
Exportação da Grade Horária semanal de um polo para .docx, no layout
oficial do modelo. Carrega o arquivo-modelo em
`app/infrastructure/templates/grade_horaria.docx` — o arquivo baixado trazia
2 blocos de tabela idênticos (um "duplicar quantos forem necessários" para
quem tem vários núcleos no mesmo documento); como aqui a grade é gerada por
polo, guardamos só 1 bloco.

Particularidades do modelo original (preservadas de propósito):
- Só existem colunas para SEGUNDA/QUARTA/SEXTA — turmas em outros dias da
  semana não aparecem na grade (limitação do documento oficial, não do
  exportador).
- A tabela tem 7 linhas de turma (A a G); a 1ª linha (A) já vem com o texto
  "SEGUNDA/QUARTA/SEXTA" nas células de valor — é sobrescrito com a carga
  horária real (ou deixado em branco) para a turma A.
- A coluna "TOTAL EM HR SEMANAIS TURMA" do modelo baixado tem uma mesclagem
  vertical quebrada (cobre da turma B em diante como uma célula só) — não é
  preenchida por não ter como fazer sentido linha a linha; os totais reais
  (carga semanal do núcleo, planejamento, total geral) vão nas 3 linhas de
  rodapé da tabela, que são células únicas e confiáveis.
- Uma legenda com a modalidade/horário de cada turma é adicionada após a
  tabela, já que o modelo só identifica as turmas por letra (A, B, C...).
"""
import io
from dataclasses import dataclass
from pathlib import Path

import docx

from app.application.relatorios.cabecalho_convenio import aplicar_cabecalho_docx

TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "infrastructure" / "templates" / "grade_horaria.docx"

DIAS_COLUNA = {"SEG": 1, "QUA": 2, "SEX": 3}
LETRAS = "ABCDEFG"
MAX_TURMAS = len(LETRAS)  # 7 — limite de linhas de turma do modelo oficial

# Linhas fixas da tabela (0-index, iguais ao modelo oficial)
LINHA_PRIMEIRA_TURMA = 2
LINHA_TOTAL_SEMANAL = 9
LINHA_PLANEJAMENTO = 10
LINHA_TOTAL_GERAL = 11
COLUNA_TOTAL = 4


@dataclass
class TurmaGrade:
    modalidade_nome: str
    horario_inicio: str  # "HH:MM"
    horario_fim: str
    dias_semana: list[str]  # ex.: ["SEG", "QUA", "SEX"]


def _duracao_horas(inicio: str, fim: str) -> float:
    hi, mi = (int(p) for p in inicio.split(":"))
    hf, mf = (int(p) for p in fim.split(":"))
    return round(((hf * 60 + mf) - (hi * 60 + mi)) / 60, 2)


def _fmt_horas(valor: float) -> str:
    if valor == int(valor):
        return str(int(valor))
    return f"{valor:.2f}".rstrip("0").rstrip(".")


def exportar_grade_horaria(
    *, polo_nome: str, turmas: list[TurmaGrade], planejamento_horas: float = 0,
    cabecalho_convenio: str | None = None,
) -> io.BytesIO:
    wb = docx.Document(TEMPLATE_PATH)
    tabela = wb.tables[0]

    tabela.cell(0, 0).text = f"HORAS AULAS SEMANAL NÚCLEO - {polo_nome}"
    # A linha da turma A já vem do modelo com "SEGUNDA/QUARTA/SEXTA" escrito
    # nas células de valor (serve de legenda para quem preenche à mão) — se
    # não houver turma nenhuma nesse polo, limpa para não sobrar esse texto
    # solto no documento gerado.
    for coluna in DIAS_COLUNA.values():
        tabela.cell(LINHA_PRIMEIRA_TURMA, coluna).text = ""

    total_semanal = 0.0
    legenda_linhas = []
    for idx, turma in enumerate(turmas[:MAX_TURMAS]):
        linha = LINHA_PRIMEIRA_TURMA + idx
        letra = LETRAS[idx]
        tabela.cell(linha, 0).text = letra

        duracao = _duracao_horas(turma.horario_inicio, turma.horario_fim)
        for dia, coluna in DIAS_COLUNA.items():
            valor = duracao if dia in turma.dias_semana else None
            tabela.cell(linha, coluna).text = _fmt_horas(valor) if valor is not None else ""
            if valor is not None:
                total_semanal += valor

        dias_label = ", ".join(d.capitalize() for d in turma.dias_semana) or "sem dia definido"
        legenda_linhas.append(
            f"{letra} — {turma.modalidade_nome} ({turma.horario_inicio}–{turma.horario_fim}, {dias_label})"
        )

    tabela.cell(LINHA_TOTAL_SEMANAL, COLUNA_TOTAL).text = f"{_fmt_horas(total_semanal)} h/semana"
    tabela.cell(LINHA_PLANEJAMENTO, COLUNA_TOTAL).text = f"{_fmt_horas(planejamento_horas)} h/semana"
    tabela.cell(LINHA_TOTAL_GERAL, COLUNA_TOTAL).text = f"{_fmt_horas(total_semanal + planejamento_horas)} h/semana"

    if legenda_linhas:
        wb.add_paragraph("")
        wb.add_paragraph("Legenda das turmas:")
        for linha_texto in legenda_linhas:
            wb.add_paragraph(f"• {linha_texto}")

    aplicar_cabecalho_docx(wb, cabecalho_convenio)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
