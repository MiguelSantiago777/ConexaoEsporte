"""Cabeçalho institucional (logo do projeto + nome do projeto/número de
convênio/vigência) aplicado a todos os relatórios exportados pelo sistema.

Estratégia: escrever a informação sempre no CABEÇALHO de página (nunca em
célula/parágrafo de conteúdo) — os modelos oficiais (Ficha de Execução,
Lista de Presença, Planilha de Núcleos, Grade Horária, Termos) têm layout
regulatório fixo (Portaria nº 102/2024); a área de cabeçalho de impressão
é sempre uma faixa separada da grade/corpo do documento, então nunca
sobrescreve nenhuma célula/parágrafo mapeado deles.

Limitação conhecida: o Excel (e o openpyxl) não suporta imagem embutida no
cabeçalho/rodapé *de impressão* de uma planilha — só texto. Por isso:
- Nos relatórios .docx (Grade Horária, Termo de Entrega, Termo de
  Responsabilidade), o logo entra de verdade no cabeçalho do documento.
- Nos .xlsx SEM layout oficial fixo (Relatório Geral, Relatório do Polo),
  o logo entra como imagem ancorada no topo da primeira área da planilha
  (ver `logo_flutuante_xlsx` / `xlsx_estilo.cabecalho_documento`).
- Nos 3 .xlsx de layout oficial fixo (Lista de Presença, Planilha de
  Núcleos, Ficha de Execução) e na exportação genérica de tabelas, só o
  texto entra no cabeçalho de impressão — não arriscamos sobrepor nenhuma
  célula do formulário oficial com uma imagem solta.
"""
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from app.domain.configuracao_geral.entities import ConfiguracaoGeral

LOGO_PATH = Path(__file__).resolve().parents[2] / "infrastructure" / "templates" / "logo.png"


def texto_cabecalho(config: ConfiguracaoGeral | None) -> str | None:
    """Monta o texto do cabeçalho a partir da Configuração Geral. Omite
    partes ausentes; retorna None se não houver nada configurado (nenhum
    cabeçalho extra é aplicado nesse caso)."""
    if not config:
        return None

    partes = []
    if config.nome_projeto:
        partes.append(f"Projeto: {config.nome_projeto}")
    if config.numero_convenio:
        partes.append(f"Convênio nº {config.numero_convenio}")
    if config.data_inicio_projeto or config.data_fim_projeto:
        inicio = config.data_inicio_projeto.strftime("%d/%m/%Y") if config.data_inicio_projeto else "?"
        fim = config.data_fim_projeto.strftime("%d/%m/%Y") if config.data_fim_projeto else "?"
        partes.append(f"Vigência do projeto: {inicio} a {fim}")

    return " — ".join(partes) if partes else None


def aplicar_cabecalho_xlsx(wb, texto: str | None, ajustar_paginacao: bool = False) -> None:
    """Texto no cabeçalho de impressão — usado em todo relatório .xlsx,
    inclusive os de layout oficial fixo (só texto, nunca imagem, aqui).

    `ajustar_paginacao=True` também força cada planilha a caber numa
    página de largura ao imprimir/exportar em PDF — só usar em planilhas
    montadas do zero por nós (Relatório Geral/Polo), nunca nos modelos
    oficiais fixos, cuja escala de impressão já vem definida (e mexer nela
    arrisca alterar como o formulário oficial fica no papel)."""
    if not texto and not ajustar_paginacao:
        return
    for ws in wb.worksheets:
        if texto:
            ws.oddHeader.center.text = texto
            ws.evenHeader.center.text = texto
            # Alguns modelos oficiais definem a margem superior quase zero
            # (pra aproveitar o máximo de área imprimível) — sem espaço, o
            # cabeçalho de impressão fica sem onde renderizar e o
            # Excel/LibreOffice simplesmente não o desenha. Garante uma
            # margem mínima pro cabeçalho aparecer, sem mexer nas margens
            # esquerda/direita (que definem o layout horizontal do
            # formulário oficial).
            margem_minima = ws.page_margins.header + 0.15
            if ws.page_margins.top < margem_minima:
                ws.page_margins.top = margem_minima
        if ajustar_paginacao:
            ws.page_setup.orientation = "landscape"
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0
            ws.sheet_properties.pageSetUpPr.fitToPage = True


def aplicar_cabecalho_docx(document, texto: str | None, com_logo: bool = True) -> None:
    """Cabeçalho do documento .docx — logo do projeto + texto do convênio,
    lado a lado e centralizados."""
    if not texto and not (com_logo and LOGO_PATH.exists()):
        return
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        for run in list(paragrafo.runs):
            run.text = ""
        if com_logo and LOGO_PATH.exists():
            paragrafo.add_run().add_picture(str(LOGO_PATH), height=Cm(1.2))
            if texto:
                paragrafo.add_run("   ")
        if texto:
            paragrafo.add_run(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER


def logo_flutuante_xlsx(ws, linha: int = 1, coluna: str = "A") -> None:
    """Ancora o logo do projeto no topo de uma planilha SEM layout oficial
    fixo (Relatório Geral/Polo) — nunca usar nos .xlsx de modelo oficial."""
    if not LOGO_PATH.exists():
        return
    from openpyxl.drawing.image import Image as ImagemXLSX

    imagem = ImagemXLSX(str(LOGO_PATH))
    imagem.height = 50
    imagem.width = 50
    ws.add_image(imagem, f"{coluna}{linha}")
