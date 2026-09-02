"""
Exportação do Termo de Responsabilidade para .docx, no layout oficial do
modelo. Carrega o arquivo-modelo em
`app/infrastructure/templates/termo_responsabilidade.docx`.

O parágrafo principal do modelo mistura o texto fixo do termo com vários
espaços em branco para preencher à mão (nome, RG, CPF, endereço, bairro,
cidade) tudo numa frase só — a forma confiável de preencher isso via
python-docx é reconstruir a frase inteira (mesma estratégia já usada nos
demais exportadores para células/parágrafos "rótulo + preenchimento
manual na mesma célula").
"""
import io
from datetime import date
from pathlib import Path

import docx

from app.application.relatorios.cabecalho_convenio import aplicar_cabecalho_docx

TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "infrastructure" / "templates" / "termo_responsabilidade.docx"

MESES_PT = [
    "", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]


def _substituir_paragrafo(paragrafo, texto: str) -> None:
    for run in paragrafo.runs:
        run.text = ""
    if paragrafo.runs:
        paragrafo.runs[0].text = texto
    else:
        paragrafo.add_run(texto)


def exportar_termo_responsabilidade(
    *,
    representante_nome: str,
    representante_rg: str,
    representante_cpf: str,
    endereco: str,
    bairro: str,
    cidade: str,
    data_assinatura: date | None = None,
    cabecalho_convenio: str | None = None,
) -> io.BytesIO:
    wb = docx.Document(TEMPLATE_PATH)
    data_assinatura = data_assinatura or date.today()

    texto_termo = (
        f"Eu, {representante_nome}, RG nº {representante_rg}, CPF nº {representante_cpf}, "
        f"residente e domiciliado(a) na {endereco}, Bairro {bairro}, Cidade {cidade}, "
        "declaro que entregarei os documentos a mim solicitados, conforme Art. 22 e seus "
        "incisos, da Portaria nº 102, de 22 de outubro de 2024, e sou responsável pelas "
        "informações nelas contidas."
    )
    texto_data = f"{cidade}, {data_assinatura.day} de {MESES_PT[data_assinatura.month]} de {data_assinatura.year}."

    for paragrafo in wb.paragraphs:
        if paragrafo.text.startswith("Eu, (nome"):
            _substituir_paragrafo(paragrafo, texto_termo)
        elif paragrafo.text.startswith("_Cidade/UF"):
            _substituir_paragrafo(paragrafo, texto_data)

    aplicar_cabecalho_docx(wb, cabecalho_convenio)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
